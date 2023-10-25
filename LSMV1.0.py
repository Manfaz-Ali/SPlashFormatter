
from PyQt6.QtWidgets import *
from PyQt6.QtGui import *
from PyQt6.QtCore import *
import time
import sys
from datetime import datetime

from PyQt6.uic import loadUiType
import docx
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
from docx.enum.style import WD_STYLE_TYPE
ui1, _ = loadUiType("MAIN4.ui")
ui2, _ = loadUiType("splash.ui")

class SplashScreen(QSplashScreen, ui2):
    def __init__(self, parent=None):
        super(QSplashScreen, self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)
        window = QMainWindow()
        window.setWindowFlag(Qt.WindowType.FramelessWindowHint)
        pixmap = QPixmap("splash.png")
        self.setPixmap(pixmap)
        

    def progress(self):
        for i in range(100):
            time.sleep(0.1)
            # self.progressBar.setValue(i)
            self.progressBar.setValue(i)


class MainApp(QMainWindow, ui1):
    def __init__(self, parent=None):
        super(MainApp, self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)

        self.doc_name = None
        self.s_grade = None
        # self.setupUi(self)
        self.doc = docx.Document()
        self.set_pageMargin()
        self.button_actions()
        self.acki = 97
        self.ackii = 97

    def button_actions(self):
        # -----------------------------------------------------------
        self.pushButton_LmDn.clicked.connect(self.save_doc)

        # --------------------------------------------------------
        self.connect_combo_box(self.comboBox_LM_To)
        self.connect_combo_box(self.comboBox_LM_SG)
        self.connect_combo_box(self.comboBox_LM_NAME)
        self.connect_combo_box(self.comboBox_LM_Rank)
        self.connect_combo_box(self.comboBox_LM_GROUP)
        self.connect_combo_box(self.comboBox_LM_TELL)
        self.connect_combo_box(self.comboBox_LM_NO)

    def save_doc(self):
        doc = self.doc
        self.upper_lm_port()
        self.midle_lm_port1()
        self.lower_lm_port()
        now = datetime.now().strftime("%H-%M-%S-%d-%m-%Y")
        doc.save(f"{now}.docx")
        self.doc = docx.Document()  # Reset doc to a new Document object
        doc = self.doc

    def connect_combo_box(self, combo_box):
        combo_box.addItems(self.load_items(combo_box))
        combo_box.currentTextChanged.connect(self.comboBox_LM_To_changed)

    def comboBox_LM_To_changed(self, text):
        combo_box = self.sender()
        if text == "Other":
            dialog = QDialog(self)
            dialog.setWindowTitle("Add Item")

            layout = QVBoxLayout(dialog)

            label = QLabel("Enter new item:", dialog)
            layout.addWidget(label)

            lineEdit = QLineEdit(dialog)
            layout.addWidget(lineEdit)

            button_layout = QVBoxLayout()
            layout.addLayout(button_layout)

            add_button = QPushButton("Add", dialog)
            button_layout.addWidget(add_button)

            close_button = QPushButton("Close", dialog)
            button_layout.addWidget(close_button)

            add_button.clicked.connect(
                lambda: self.add_item(dialog, combo_box, lineEdit.text()))
            combo_box.addItem(lineEdit.text())
            combo_box.setCurrentIndex(combo_box.count() - 1)
            close_button.clicked.connect(dialog.close)

            dialog.exec()

    def add_item(self, dialog, combo_box, text):
        index = combo_box.findText(text)
        if index == -1:
            combo_box.addItem(text)
            combo_box.setCurrentIndex(combo_box.count() - 1)

        self.save_items(combo_box)
        dialog.close()

    def load_items(self, combo_box):
        try:
            with open(f"{combo_box.objectName()}.txt", "r") as f:
                items = f.read().splitlines()
                items = set(items)
        except:
            items = ["Item 1", "Item 2", "Item 3", "Other"]
        return items

    def save_items(self, combo_box):
        with open(f"{combo_box.objectName()}.txt", "w") as f:
            for i in range(combo_box.count()):
                f.write(combo_box.itemText(i) + "\n")

    def reset1(self):
        self.acki = 97

    def reset2(self):
        self.ackii = 97

    def get_LmRef(self):
        ab = self.comboBox_LM_NO.currentText()
        return ab

    def get_LmDate(self):
        selected_date = self.calendarWidget_LM.selectedDate()
        return selected_date

    def get_To(self):
        ab = self.comboBox_LM_To.currentText()
        return ab

    def get_LmSGRD(self):
        ab = self.comboBox_LM_SG.currentText()
        return ab

    def get_LmForName(self):
        ab = self.comboBox_LM_NAME.currentText()
        return ab

    def get_LmForRank(self):
        ab = self.comboBox_LM_Rank.currentText()
        return ab

    def get_GrpOfLmFor(self):
        ab = self.comboBox_LM_GROUP.currentText()
        return ab

    def get_tellOfLmFor(self):
        ab = self.comboBox_LM_TELL.currentText()
        return ab

    def get_LmSubject(self):
        sb = self.lineEdit_lmSubject.text()
        return sb

    def Lm_paragraph_text1(self):
        para = self.plainTextEdit_LmParagraph.toPlainText().strip()
        return para

    def Lm_paragraph_text2(self):
        para = self.plainTextEdit_LmParagraph_2.toPlainText().strip()
        return para

    def Lm_paragraph_text3(self):
        para = self.plainTextEdit_LmParagraph_3.toPlainText().strip()
        return para

    def Lm_paragraph_text4(self):
        para = self.plainTextEdit_LmParagraph_4.toPlainText().strip()
        return para

    def Lm_Sub_paragraph_text(self):
        para = self.plainTextEdit_LmSParagraph.toPlainText().strip()
        return para

    def LmHeader(self, headerText):
        doc = self.doc
        header = doc.sections[0].header
        paragraph = header.paragraphs[0]
        paragraph.add_run(headerText)
        # set the paragraph properties
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = docx.shared.Pt(12)
        paragraph.style.font.bold = False

    def LmFooter(self, footerText):
        doc = self.doc
        footer = doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.add_run(footerText)
        # set the paragraph properties
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = docx.shared.Pt(12)
        paragraph.style.font.bold = False

    def LmHdrFtr(self, grade):
        self.LmHeader(grade)
        self.LmFooter(grade)

    def para_space_handler(self):
        doc = self.doc
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.0

        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.space_after = 0

    def upper_lm_port(self):
        sg = self.get_LmSGRD()
        self.LmHdrFtr(sg)
        to = self.get_To()
        To = to.upper()
        sub = self.get_LmSubject()
        SUB = sub.upper()
        doc = self.doc
        doc.add_paragraph("")
        doc.add_paragraph("")
        org_name = "Avionics Production Factory".upper()
        title = doc.add_paragraph(org_name)
        title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        title.style.font.bold = True
        title.style.font.name = "Arial"
        title.style.font.size = docx.shared.Pt(12)
        run = title.runs[0]
        run.font.bold = True

        sub_title = "(DDD)"
        stitle = doc.add_paragraph(sub_title)
        stitle.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        stitle.style.font.bold = False
        stitle.style.font.name = "Arial"
        stitle.style.font.size = docx.shared.Pt(12)

        doc.add_paragraph("")
        doc.add_paragraph("")

        ref_to = doc.add_paragraph(To)
        ref_to.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph("")
        doc.add_paragraph("")

        if SUB == '':
            sub = "Subject".upper()
        else:
            sub = SUB.upper()
        start = doc.add_paragraph(sub)
        start.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        start.style.font.name = "Arial"
        run = start.runs[0]
        run.font.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph("")
        self.para_space_handler()

    def midle_lm_port1(self):
        doc = self.doc
        doc.add_paragraph("")
        your_para1 = self.Lm_paragraph_text1()
        your_para2 = self.Lm_paragraph_text2()
        your_para3 = self.Lm_paragraph_text3()
        your_para4 = self.Lm_paragraph_text4()

        pragraphs = [your_para1, your_para2, your_para3, your_para4]
        for your_para in pragraphs:
            if your_para == '':
                continue  # skip empty paragraphs
            else:
                paragraph = your_para

            paragraph = '\t' + paragraph
            paragraph = doc.add_paragraph(paragraph)
            paragraph.style = 'List Number'
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)
            font.bold = False

        self.para_space_handler()

    def midle_lm_port2(self):
        doc = self.doc
        ascii_val = self.acki
        char_val = chr(ascii_val)
        doc.add_paragraph("")
        paragraph = self.Lm_Sub_paragraph_text()
        paragraph = '\t'+f'({char_val})' + paragraph
        paragraph = doc.add_paragraph(paragraph)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = False
        self.para_space_handler()
        self.acki += 1

    def stamp_maker(self, NAME, RANK, GROUP, TEL):
        doc = self.doc

        # determine maximum length of each field
        max_name_len = 20
        max_rank_len = 20
        max_group_len = 20
        max_tel_len = 20

        # add a left indent of 7 steps
        left_indent = 24 * docx.shared.Pt(12)

        # left justify the name field with left indent
        name = doc.add_paragraph()
        name.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        name.paragraph_format.left_indent = left_indent
        name.add_run(f"{NAME}".ljust(max_name_len))
        name.style.font.name = "Arial"
        name.style.font.size = Pt(12)
        name.style.font.bold = True

        # left justify the rank field with left indent
        rank = doc.add_paragraph()
        rank.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        rank.paragraph_format.left_indent = left_indent
        rank.add_run(f"{RANK}".ljust(max_rank_len))
        rank.style.font.name = "Arial"
        rank.style.font.size = Pt(12)
        rank.style.font.bold = False

        # left justify the group field with left indent
        group = doc.add_paragraph()
        group.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        group.paragraph_format.left_indent = left_indent
        group.add_run(f"{GROUP}".ljust(max_group_len))
        group.style.font.name = "Arial"
        group.style.font.size = Pt(12)
        group.style.font.bold = False

        # left justify the tel field with left indent
        tel = doc.add_paragraph()
        tel.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        tel.paragraph_format.left_indent = left_indent
        tel.add_run(f"Tel Ext: {TEL}".ljust(max_tel_len))
        tel.style.font.name = "Arial"
        tel.style.font.size = Pt(12)
        tel.style.font.bold = False

    def lower_lm_port(self):
        doc = self.doc
        for i in range(5):
            doc.add_paragraph("")
        lm_no = self.get_LmRef()
        date = self.get_LmDate()
        nam = self.get_LmForName()
        rnk = self.get_LmForRank()
        grp = self.get_GrpOfLmFor()
        tel = self.get_tellOfLmFor()
        formatted_date = date.toString("dd MMMM, yyyy")
        LM_REFRENCE = lm_no.upper()
        # DATE = date.upper()
        NAME = nam.upper()
        RANK = rnk.upper()
        GROUP = grp.upper()
        TEL = tel.upper()
        self.stamp_maker(NAME, RANK, GROUP, TEL)
        doc.add_paragraph("")
        lm_ref = doc.add_paragraph()
        lm_ref.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        lm_ref.style.font.bold = False
        lm_ref.style.font.name = "Arial"

        lm_ref.add_run("LM No  ")
        lm_ref.add_run(LM_REFRENCE)
        lm_ref.add_run(" ")
        lm_ref.add_run("   dated    ")
        lm_ref.add_run(formatted_date)

        self.para_space_handler()

    def set_pageMargin(self):
        doc = self.doc
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(0.5)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    splash = SplashScreen()
    splash.show()
    splash.progress()
    window = MainApp()
    window.show()
    splash.finish(window)
    app.exec()
